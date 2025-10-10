"""
DOCX parser implementation for Study Buddy MCP Server.

This module implements Microsoft Word document parsing using python-docx,
following the Strategy pattern and Clean Architecture Layer 4 principles.
"""

import logging
import os
from typing import Any, Dict

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    raise ImportError(
        "python-docx is required for DOCX parsing. Install with: pip install python-docx"
    )

from ..models.parse_result import ParseResult

from .base_parser import BaseParser, ParseError


class DOCXParser(BaseParser):
    """
    Microsoft Word DOCX document parser using python-docx.

    This class implements the BaseParser interface for DOCX documents,
    providing comprehensive content and metadata extraction.

    Features:
    - Full text extraction from paragraphs and tables
    - Document metadata extraction (title, author, etc.)
    - Style and formatting information
    - Table content extraction
    - Header and footer content
    - Document statistics

    Clean Architecture Layer 4: Infrastructure Implementation
    - Implements parser strategy interface
    - No dependencies on business logic
    - Pure document processing logic
    """

    def __init__(self):
        """Initialize DOCX parser with logging."""
        self.logger = logging.getLogger(__name__)

    def supports_file_type(self, file_path: str) -> bool:
        """
        Check if file is a DOCX document.

        Args:
            file_path: Path to file to check

        Returns:
            True if file has .docx extension
        """
        return file_path.lower().endswith(".docx")

    def get_supported_extensions(self) -> list[str]:
        """Get supported file extensions."""
        return ["docx"]

    def parse(self, file_path: str) -> ParseResult:
        """
        Parse DOCX document and extract content with metadata.

        Args:
            file_path: Absolute path to DOCX file

        Returns:
            ParseResult with extracted content and metadata

        Raises:
            ParseError: If DOCX parsing fails
            FileNotFoundError: If file doesn't exist
        """
        # Validate file first
        self.validate_file(file_path)

        try:
            # Open DOCX document
            doc = Document(file_path)

            # Extract content
            content = self._extract_content(doc)

            if not content.strip():
                raise ParseError(
                    "No text content could be extracted from DOCX document",
                    file_path=file_path,
                )

            # Extract metadata
            metadata = self._extract_metadata(doc, file_path)

            return ParseResult(content=content, metadata=metadata)

        except FileNotFoundError:
            raise
        except ParseError:
            raise
        except PackageNotFoundError as e:
            raise ParseError(
                "Invalid or corrupted DOCX file",
                file_path=file_path,
                original_error=e,
            )
        except Exception as e:
            self.logger.error(
                f"Unexpected error parsing DOCX {file_path}: {e}"
            )
            raise ParseError(
                "Failed to parse DOCX document",
                file_path=file_path,
                original_error=e,
            )

    def _extract_content(self, doc) -> str:
        """
        Extract all text content from DOCX document.

        Args:
            doc: python-docx Document object

        Returns:
            Extracted text content
        """
        content_parts = []

        try:
            # Extract paragraph content
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # Extract table content
            for table in doc.tables:
                table_content = self._extract_table_content(table)
                if table_content:
                    content_parts.append(table_content)

            # Extract header content
            for section in doc.sections:
                # Header content
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content_parts.append(f"[HEADER] {text}")

                # Footer content
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content_parts.append(f"[FOOTER] {text}")

            # Join all content
            full_content = "\n\n".join(content_parts)

            self.logger.info(
                f"Extracted {len(full_content)} characters from DOCX"
            )
            return full_content

        except Exception as e:
            self.logger.error(f"Error extracting DOCX content: {e}")
            raise ParseError(f"Failed to extract content from DOCX: {e}")

    def _extract_table_content(self, table) -> str:
        """
        Extract content from a table.

        Args:
            table: python-docx Table object

        Returns:
            Formatted table content
        """
        try:
            table_rows = []

            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_cells.append(cell_text or "")

                if any(cell for cell in row_cells):  # Skip empty rows
                    table_rows.append(" | ".join(row_cells))

            if table_rows:
                return "\n".join(table_rows)

            return ""

        except Exception as e:
            self.logger.warning(f"Error extracting table content: {e}")
            return ""

    def _extract_metadata(self, doc, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: python-docx Document object
            file_path: Path to DOCX file

        Returns:
            Dictionary containing DOCX metadata
        """
        metadata = {
            "file_type": "docx",
            "file_size": os.path.getsize(file_path),
            "parser": self.get_parser_name(),
        }

        try:
            # Document properties
            core_props = doc.core_properties

            if core_props.title:
                metadata["title"] = core_props.title

            if core_props.author:
                metadata["author"] = core_props.author

            if core_props.subject:
                metadata["subject"] = core_props.subject

            if core_props.keywords:
                metadata["keywords"] = core_props.keywords

            if core_props.comments:
                metadata["comments"] = core_props.comments

            if core_props.category:
                metadata["category"] = core_props.category

            # Dates
            if core_props.created:
                metadata["created_date"] = core_props.created.isoformat()

            if core_props.modified:
                metadata["modified_date"] = core_props.modified.isoformat()

            if core_props.last_printed:
                metadata["last_printed"] = core_props.last_printed.isoformat()

            # Document statistics
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            metadata["section_count"] = len(doc.sections)

            # Style information
            styles = [style.name for style in doc.styles if style.name]
            metadata["available_styles"] = styles[:20]  # Limit to first 20

            # Page setup information
            if doc.sections:
                first_section = doc.sections[0]
                page_setup = {
                    "page_width": str(first_section.page_width),
                    "page_height": str(first_section.page_height),
                    "left_margin": str(first_section.left_margin),
                    "right_margin": str(first_section.right_margin),
                    "top_margin": str(first_section.top_margin),
                    "bottom_margin": str(first_section.bottom_margin),
                }
                metadata["page_setup"] = page_setup

        except Exception as e:
            self.logger.warning(f"Could not extract all DOCX metadata: {e}")
            # Continue without full metadata - not critical

        return metadata
