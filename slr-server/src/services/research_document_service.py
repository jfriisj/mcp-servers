"""
Research Document Service for academic paper lifecycle orchestration.

This module implements the ResearchDocumentService class following Clean Architecture
Layer 2 principles, orchestrating academic document operations while remaining
framework-agnostic and focusing on systematic literature review workflows.
"""

import os
import re
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path

from ..domain.models import ResearchPaper, Author, Journal
from ..repositories.paper_repository import PaperRepository

logger = logging.getLogger(__name__)


class ResearchDocumentService:
    """
    Research document business logic service for systematic literature reviews.

    Orchestrates academic paper operations including upload, metadata extraction,
    citation analysis, paper classification, and research corpus management while
    enforcing academic research business rules and coordinating between repositories.

    Follows Clean Architecture Layer 2 principles:
    - Framework-agnostic business logic
    - Depends on abstractions (repositories)
    - Contains reusable academic research rules
    - Validates academic input and enforces constraints
    - Coordinates operations across academic components

    SOLID Principles:
    - SRP: Single responsibility for research document business logic
    - OCP: Open for extension via academic strategy patterns
    - LSP: Can be substituted with other research document services
    - ISP: Focused interface for academic document operations only
    - DIP: Depends on repository abstractions

    Academic Research Focus:
    - PRISMA compliance for systematic reviews
    - Academic metadata extraction and validation
    - Citation network analysis and management
    - Research corpus organization and filtering
    - Quality assessment integration
    """

    # Academic business rule constants
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB limit for academic papers
    ACADEMIC_EXTENSIONS = {'.pdf', '.docx', '.tex', '.bib', '.ris'}
    MIN_ABSTRACT_LENGTH = 50  # Minimum abstract length for quality papers
    MAX_AUTHORS = 50  # Reasonable limit for author list
    VALID_STUDY_TYPES = {'experimental', 'observational', 'review', 'meta-analysis', 'case_study', 'survey'}
    VALID_METHODOLOGIES = {'quantitative', 'qualitative', 'mixed_methods', 'theoretical', 'empirical'}

    def __init__(self, paper_repository: PaperRepository):
        """
        Initialize ResearchDocumentService with required dependencies.

        Args:
            paper_repository: Repository for research paper persistence

        Note:
            Dependencies are injected to enable testing and flexibility.
        """
        self.paper_repository = paper_repository

    # ===== Helper Methods for upload_paper (Extracted for SRP) =====

    def _validate_file_path(self, file_path: str) -> Tuple[str, int, int]:
        """
        Validate file exists, is accessible, and meets academic standards.
        
        Returns:
            Tuple of (file_extension, file_size, file_hash)
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file validation fails
        """
        # Academic Rule: File must exist and be accessible
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Academic paper file not found: {file_path}")

        if not os.path.isfile(file_path):
            raise ValueError(f"Path is not a file: {file_path}")

        # Academic Rule: File size appropriate for academic papers
        file_size = os.path.getsize(file_path)
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(
                f"Academic paper size {file_size} bytes exceeds maximum of {self.MAX_FILE_SIZE} bytes"
            )

        # Academic Rule: Must be academic document format
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.ACADEMIC_EXTENSIONS:
            raise ValueError(
                f"Unsupported academic format '{file_ext}'. Supported: {', '.join(self.ACADEMIC_EXTENSIONS)}"
            )

        # Academic Rule: No duplicate file paths
        if self.paper_repository.get_by_file_path(file_path):
            raise ValueError(f"Research paper already exists for file path: {file_path}")

        return file_ext, file_size, 0  # hash placeholder

    def _extract_and_merge_metadata(
        self,
        file_path: str,
        auto_extract: bool,
        title: Optional[str],
        authors: Optional[List[Author]],
        journal: Optional[Journal],
        publication_year: Optional[int],
        doi: Optional[str]
    ) -> Dict[str, Any]:
        """
        Extract metadata from file and merge with provided values.
        
        Provided values take precedence over extracted values.
        
        Returns:
            Merged metadata dictionary
        """
        # Extract academic metadata if requested
        extracted_metadata = {}
        if auto_extract:
            try:
                extracted_metadata = self.extract_metadata(file_path)
            except Exception as e:
                raise ResearchDocumentError(f"Failed to extract academic metadata: {str(e)}") from e

        # Combine provided and extracted metadata with precedence to provided
        return {
            "title": title or extracted_metadata.get("title") or Path(file_path).stem,
            "authors": authors or extracted_metadata.get("authors", []),
            "journal": journal or extracted_metadata.get("journal"),
            "publication_year": publication_year or extracted_metadata.get("publication_year"),
            "abstract": extracted_metadata.get("abstract", ""),
            "keywords": extracted_metadata.get("keywords", []),
            "sample_size": extracted_metadata.get("sample_size"),
            "total_pages": extracted_metadata.get("total_pages"),
            "total_words": extracted_metadata.get("total_words"),
        }

    def _validate_paper_metadata(
        self,
        title: str,
        authors: List[Author],
        publication_year: Optional[int],
        abstract: Optional[str],
        doi: Optional[str]
    ) -> None:
        """
        Validate paper metadata meets academic standards.
        
        Raises:
            ValueError: If validation fails
        """
        # Academic Rule: Title must be meaningful
        if not title.strip():
            raise ValueError("Research paper title cannot be empty")

        # Academic Rule: Author count within limits
        if len(authors) > self.MAX_AUTHORS:
            raise ValueError(f"Author count ({len(authors)}) exceeds maximum of {self.MAX_AUTHORS}")

        # Academic Rule: Publication year validation
        current_year = datetime.now().year
        if publication_year and (publication_year < 1900 or publication_year > current_year + 1):
            raise ValueError(f"Invalid publication year: {publication_year}")

        # Academic Rule: Abstract quality standards
        if abstract and len(abstract) < self.MIN_ABSTRACT_LENGTH:
            raise ValueError(
                f"Abstract too short ({len(abstract)} chars). Minimum: {self.MIN_ABSTRACT_LENGTH}"
            )

        # Academic Rule: DOI uniqueness if provided
        if doi and self._is_duplicate_doi(doi):
            raise ValueError(f"Research paper with DOI '{doi}' already exists")

    def _build_research_paper_entity(
        self,
        file_path: str,
        file_ext: str,
        file_size: int,
        metadata: Dict[str, Any],
        doi: Optional[str],
        tags: Optional[List[str]]
    ) -> ResearchPaper:
        """
        Build ResearchPaper entity from validated metadata.
        
        Returns:
            Constructed ResearchPaper object
        """
        # Classify paper academic characteristics
        classification = self._classify_paper(
            metadata["title"],
            metadata["abstract"],
            metadata["keywords"]
        )

        # Create research paper entity
        return ResearchPaper(
            title=metadata["title"].strip(),
            file_path=file_path,
            file_type=file_ext[1:],
            authors=metadata["authors"],
            journal=metadata["journal"],
            publication_year=metadata["publication_year"],
            doi=doi,
            abstract=metadata["abstract"],
            keywords=metadata["keywords"],
            methodology=classification.get("methodology"),
            study_type=classification.get("study_type"),
            sample_size=metadata.get("sample_size"),
            citation_count=0,  # Will be updated by citation analysis
            upload_date=datetime.now(),
            file_size=file_size,
            total_pages=metadata.get("total_pages"),
            total_words=metadata.get("total_words"),
            tags=tags or [],
            indexed=False,
            quality_assessed=False,
            included_in_review=None,  # To be determined by quality assessment
            notes=""
        )

    # ===== Main upload_paper method (Refactored for SRP) =====

    def upload_paper(
        self,
        file_path: str,
        title: Optional[str] = None,
        authors: Optional[List[Author]] = None,
        journal: Optional[Journal] = None,
        publication_year: Optional[int] = None,
        doi: Optional[str] = None,
        tags: Optional[List[str]] = None,
        auto_extract_metadata: bool = True
    ) -> ResearchPaper:
        """
        Upload and process a new research paper with academic validation.

        Academic Business Logic:
        1. Validate file meets academic standards
        2. Extract academic metadata (authors, citations, abstract)
        3. Classify paper methodology and study type
        4. Perform citation analysis
        5. Create research paper entity with academic attributes
        6. Persist in research corpus

        Args:
            file_path: Absolute path to academic paper file
            title: Optional title override
            authors: Optional author list
            journal: Optional journal information
            publication_year: Year of publication
            doi: Digital Object Identifier
            tags: Research topic tags
            auto_extract_metadata: Whether to extract metadata from file

        Returns:
            Created ResearchPaper with populated academic metadata

        Raises:
            ValueError: If file validation fails or academic rules violated
            FileNotFoundError: If file doesn't exist
            ResearchDocumentError: If processing or persistence fails

        Academic Rules Enforced:
        - File must be academic format (PDF, DOCX, TEX, BIB)
        - File size appropriate for academic papers
        - No duplicate DOIs or file paths
        - Abstract meets minimum quality standards
        - Author count within reasonable limits
        - Publication year is valid academic year
        """
        # 1. Validate file
        file_ext, file_size, file_path_obj = self._validate_file_path(file_path)

        # 2. Extract and merge metadata
        metadata = self._extract_and_merge_metadata(
            file_path=file_path,
            auto_extract=auto_extract_metadata,
            title=title,
            authors=authors,
            journal=journal,
            publication_year=publication_year,
            doi=doi
        )

        # 3. Validate academic metadata
        self._validate_paper_metadata(
            title=metadata["title"],
            authors=metadata["authors"],
            publication_year=metadata["publication_year"],
            abstract=metadata.get("abstract"),
            doi=doi
        )

        # 4. Build entity
        paper = self._build_research_paper_entity(
            file_path=file_path,
            file_ext=file_ext,
            file_size=file_size,
            metadata=metadata,
            doi=doi,
            tags=tags
        )

        # 5. Persist and analyze
        try:
            created_paper = self.paper_repository.create(paper)
            
            # Perform citation analysis after creation
            if auto_extract_metadata and created_paper.id:
                self._analyze_paper_citations(created_paper.id, file_path)
            
            return created_paper

        except Exception as e:
            raise ResearchDocumentError(f"Failed to create research paper: {str(e)}") from e

    def upload_paper_with_full_text(
        self,
        file_path: str,
        title: Optional[str] = None,
        authors: Optional[List[Author]] = None,
        journal: Optional[Journal] = None,
        publication_year: Optional[int] = None,
        doi: Optional[str] = None,
        tags: Optional[List[str]] = None,
        auto_extract_metadata: bool = True,
        replace_existing: bool = True
    ) -> Tuple[ResearchPaper, bool]:
        """
        Upload research paper with full text, replacing existing if needed.

        This method handles full-text paper uploads with a priority flag to ensure
        full-text versions are weighted higher than versions without full text.

        Academic Business Logic:
        1. Check if paper already exists (by title, DOI, or file path)
        2. If exists and replace_existing=True: Update with full-text version
        3. If exists and replace_existing=False: Return existing without update
        4. If new: Create new paper record

        Args:
            file_path: Absolute path to academic paper file (full-text version)
            title: Optional title override
            authors: Optional author list
            journal: Optional journal information
            publication_year: Year of publication
            doi: Digital Object Identifier (used for duplicate detection)
            tags: Research topic tags including 'full-text' tag
            auto_extract_metadata: Whether to extract metadata from file
            replace_existing: Whether to replace existing papers with full-text version

        Returns:
            Tuple of (ResearchPaper, is_new_upload)
            - is_new_upload: True if new paper, False if existing paper updated/returned

        Raises:
            ValueError: If file validation fails or academic rules violated
            FileNotFoundError: If file doesn't exist
            ResearchDocumentError: If processing or persistence fails

        Academic Rules:
        - Full-text versions are prioritized over abstracts-only
        - Updates preserve existing screening decisions and metadata
        - Proper tagging to identify full-text papers
        """
        # 1. Validate file
        file_ext, file_size, file_path_obj = self._validate_file_path(file_path)

        # 2. Extract and merge metadata
        metadata = self._extract_and_merge_metadata(
            file_path=file_path,
            auto_extract=auto_extract_metadata,
            title=title,
            authors=authors,
            journal=journal,
            publication_year=publication_year,
            doi=doi
        )

        # 3. Validate academic metadata
        self._validate_paper_metadata(
            title=metadata["title"],
            authors=metadata["authors"],
            publication_year=metadata["publication_year"],
            abstract=metadata.get("abstract"),
            doi=doi
        )

        # 4. Check for existing paper
        existing_paper = None
        
        # Try to find by DOI
        if doi:
            try:
                existing = self.paper_repository.get_by_doi(doi)
                if existing:
                    existing_paper = existing
            except Exception as e:
                logger.debug(f"Could not find paper by DOI: {e}")
        
        # Try to find by title similarity
        if not existing_paper and metadata.get("title"):
            try:
                # Simple title matching
                all_papers = self.paper_repository.list_all()
                for paper in all_papers:
                    if paper.title and paper.title.lower() == metadata["title"].lower():
                        existing_paper = paper
                        break
            except Exception:
                # Ignore errors during title matching
                pass

        # 5. Handle update or create
        if existing_paper:
            if replace_existing:
                # Update existing paper with full-text version
                existing_paper.abstract = metadata.get("abstract", existing_paper.abstract)
                existing_paper.total_pages = metadata.get("total_pages", existing_paper.total_pages)
                existing_paper.total_words = metadata.get("total_words", existing_paper.total_words)
                existing_paper.file_path = file_path
                existing_paper.file_type = file_ext[1:]
                existing_paper.file_size = file_size
                
                # Add full-text tag if not present
                if existing_paper.tags:
                    if "full-text" not in existing_paper.tags:
                        existing_paper.tags.append("full-text")
                else:
                    existing_paper.tags = ["full-text"]
                
                # Merge additional tags
                if tags:
                    for tag in tags:
                        if tag not in existing_paper.tags:
                            existing_paper.tags.append(tag)
                
                try:
                    updated_paper = self.paper_repository.update(existing_paper)
                    return (updated_paper, False)  # False indicates update, not new
                except Exception as e:
                    raise ResearchDocumentError(
                        f"Failed to update research paper with full text: {str(e)}"
                    ) from e
            else:
                # Return existing without update
                return (existing_paper, False)

        # 6. Create new paper if no existing version found
        paper = self._build_research_paper_entity(
            file_path=file_path,
            file_ext=file_ext,
            file_size=file_size,
            metadata=metadata,
            doi=doi,
            tags=tags if tags else ["full-text"]
        )

        # Ensure full-text tag is present
        if "full-text" not in paper.tags:
            paper.tags.append("full-text")

        # 7. Persist and analyze
        try:
            created_paper = self.paper_repository.create(paper)
            
            # Perform citation analysis after creation
            if auto_extract_metadata and created_paper.id:
                self._analyze_paper_citations(created_paper.id, file_path)
            
            return (created_paper, True)  # True indicates new upload

        except Exception as e:
            raise ResearchDocumentError(f"Failed to create research paper: {str(e)}") from e

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract academic metadata from research paper file.

        Academic metadata includes:
        - Title and abstract
        - Author information
        - Journal and publication details
        - Keywords and research topics
        - Citations and references
        - Document statistics

        Args:
            file_path: Path to academic paper file

        Returns:
            Dictionary containing extracted academic metadata

        Raises:
            ResearchDocumentError: If metadata extraction fails
        """
        try:
            # Basic file information
            file_stats = os.stat(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            metadata = {
                "file_type": file_ext[1:],
                "file_size": file_stats.st_size,
                "created_date": datetime.fromtimestamp(file_stats.st_ctime),
                "modified_date": datetime.fromtimestamp(file_stats.st_mtime)
            }

            # Extract content-based metadata based on file type
            if file_ext == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_ext == '.docx':
                metadata.update(self._extract_docx_metadata(file_path))
            elif file_ext == '.tex':
                metadata.update(self._extract_tex_metadata(file_path))
            elif file_ext == '.bib':
                metadata.update(self._extract_bib_metadata(file_path))
            elif file_ext == '.ris':
                metadata.update(self._extract_ris_metadata(file_path))
            else:
                # Fallback to basic text extraction
                metadata.update(self._extract_text_metadata(file_path))

            # Enhance metadata with academic analysis
            metadata.update(self._enhance_academic_metadata(metadata))

            return metadata

        except Exception as e:
            raise ResearchDocumentError(f"Failed to extract metadata: {str(e)}") from e

    def analyze_citations(self, paper_id: int) -> Dict[str, Any]:
        """
        Perform comprehensive citation analysis for a research paper.

        Academic Citation Analysis:
        - Extract cited references
        - Identify citation patterns and networks
        - Calculate citation metrics
        - Detect self-citations and citation clusters
        - Analyze temporal citation patterns

        Args:
            paper_id: ID of research paper to analyze

        Returns:
            Dictionary containing citation analysis results

        Raises:
            ResearchDocumentError: If citation analysis fails
        """
        paper = self.paper_repository.get_by_id(paper_id)
        if not paper:
            raise ValueError(f"Research paper {paper_id} not found")

        try:
            return self._analyze_paper_citations(paper_id, paper.file_path)
        except Exception as e:
            raise ResearchDocumentError(f"Citation analysis failed: {str(e)}") from e

    def classify_paper(self, paper_id: int) -> Dict[str, Any]:
        """
        Classify research paper by methodology and study type.

        Academic Classification:
        - Methodology: quantitative, qualitative, mixed methods, theoretical, empirical
        - Study type: experimental, observational, review, meta-analysis, case study, survey
        - Research domain and topic classification
        - Quality indicators and research rigor assessment

        Args:
            paper_id: ID of research paper to classify

        Returns:
            Dictionary containing classification results

        Raises:
            ResearchDocumentError: If classification fails
        """
        paper = self.paper_repository.get_by_id(paper_id)
        if not paper:
            raise ValueError(f"Research paper {paper_id} not found")

        try:
            classification = self._classify_paper(paper.title, paper.abstract or "", paper.keywords)
            
            # Update paper with classification
            if classification.get("methodology") and not paper.methodology:
                paper.methodology = classification["methodology"]
            if classification.get("study_type") and not paper.study_type:
                paper.study_type = classification["study_type"]
                
            self.paper_repository.update(paper)
            
            return classification

        except Exception as e:
            raise ResearchDocumentError(f"Paper classification failed: {str(e)}") from e

    def get_research_corpus(
        self,
        filters: Optional[Dict[str, Any]] = None,
        sort_by: str = "publication_year",
        sort_order: str = "desc",
        limit: Optional[int] = None
    ) -> List[ResearchPaper]:
        """
        Retrieve and manage research corpus with academic filtering.

        Academic Corpus Management:
        - Filter by research methodology and study type
        - Filter by publication period and venue
        - Filter by author, journal, and citation metrics
        - Support systematic review inclusion/exclusion criteria
        - Quality assessment status filtering

        Args:
            filters: Academic filter criteria
            sort_by: Sort field (publication_year, citation_count, title)
            sort_order: Sort order (asc, desc)
            limit: Maximum papers to return

        Returns:
            List of research papers matching academic criteria

        Raises:
            ValueError: If filter parameters are invalid
        """
        # Academic Rule: Validate sort parameters
        valid_sort_fields = {"publication_year", "citation_count", "title", "created_at", "upload_date"}
        if sort_by not in valid_sort_fields:
            raise ValueError(f"Invalid sort field '{sort_by}'. Valid: {', '.join(valid_sort_fields)}")

        if sort_order.lower() not in {"asc", "desc"}:
            raise ValueError("Sort order must be 'asc' or 'desc'")

        # Academic Rule: Reasonable corpus size limits
        if limit is not None and limit > 10000:
            raise ValueError("Corpus limit cannot exceed 10,000 papers")

        try:
            # Apply academic filters
            academic_filters = self._build_academic_filters(filters or {})
            
            # Get filtered papers
            papers = self.paper_repository.list_all(academic_filters)
            
            # Apply sorting
            papers = self._sort_research_papers(papers, sort_by, sort_order)
            
            # Apply limit
            if limit is not None:
                papers = papers[:limit]
            
            return papers

        except Exception as e:
            raise ResearchDocumentError(f"Failed to retrieve research corpus: {str(e)}") from e
    
    def process_document(
        self,
        file_path: str,
        title: Optional[str] = None,
        authors: Optional[List[str]] = None,
        publication_year: Optional[int] = None,
        doi: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> ResearchPaper:
        """
        Process document upload (adapter method for MCP interface).
        
        This method adapts the MCP interface to the upload_paper business logic.
        
        Args:
            file_path: Path to document file
            title: Optional title override
            authors: List of author names (will be converted to Author objects)
            publication_year: Year of publication
            doi: Digital Object Identifier
            tags: Research topic tags
        
        Returns:
            Created ResearchPaper instance
        
        Raises:
            ResearchDocumentError: If processing fails
        """
        try:
            # Convert author names to Author objects
            author_objects = []
            if authors:
                for author_name in authors:
                    from ..domain.models import Author
                    author_objects.append(Author(name=author_name))
            
            # Call the main upload_paper method
            return self.upload_paper(
                file_path=file_path,
                title=title,
                authors=author_objects,
                publication_year=publication_year,
                doi=doi,
                tags=tags or []
            )
            
        except Exception as e:
            raise ResearchDocumentError(f"Failed to process document: {str(e)}") from e
    
    def search_documents(
        self,
        query: str,
        search_type: str = "semantic",
        filters: Optional[Dict[str, Any]] = None,
        limit: int = 20
    ) -> List[ResearchPaper]:
        """
        Search documents (adapter method for MCP interface).
        
        This method adapts the MCP search interface to the paper search business logic.
        
        Args:
            query: Search query string
            search_type: Type of search (semantic, keyword, citation)
            filters: Optional search filters
            limit: Maximum results to return
        
        Returns:
            List of matching ResearchPaper instances
        
        Raises:
            ResearchDocumentError: If search fails
        """
        try:
            # For now, use the existing search_papers method
            # In future versions, different search_type values could trigger different search strategies
            return self.search_papers(query, limit=limit)
            
        except Exception as e:
            raise ResearchDocumentError(f"Failed to search documents: {str(e)}") from e

    def search_papers(
        self,
        query: str,
        search_fields: Optional[List[str]] = None,
        filters: Optional[Dict[str, Any]] = None,
        limit: int = 20
    ) -> List[ResearchPaper]:
        """
        Academic search across research papers with field-specific searching.

        Academic Search Features:
        - Full-text search across title, abstract, and content
        - Author and journal-specific searches
        - Citation-aware searching
        - Research methodology and topic searches
        - Boolean and phrase search support

        Args:
            query: Academic search query
            search_fields: Specific fields to search (title, abstract, authors, keywords)
            filters: Additional academic filters
            limit: Maximum results to return

        Returns:
            List of matching research papers ordered by relevance

        Raises:
            ValueError: If search parameters are invalid
        """
        # Academic Rule: Valid search query
        if not query or not query.strip():
            raise ValueError("Academic search query cannot be empty")

        query = query.strip()
        if len(query) > 1000:
            raise ValueError("Search query cannot exceed 1000 characters")

        # Academic Rule: Reasonable search limits
        if limit > 200:
            raise ValueError("Search limit cannot exceed 200 results")

        try:
            # If search fields specified, perform targeted search
            if search_fields:
                return self._targeted_academic_search(query, search_fields, filters, limit)
            else:
                # Perform full-text academic search
                return self.paper_repository.search_papers(query, limit=limit)
                
        except Exception as e:
            raise ResearchDocumentError(f"Academic search failed: {str(e)}") from e
    
    def _extract_authors_from_text(self, text: str) -> List['Author']:
        """Extract author names from text content."""
        from ..domain.models import Author
        authors = []
        
        # Common patterns for author extraction
        patterns = [
            r'(?:Authors?|By):?\s*([A-Z][^\n]*?)(?:\n|$)',
            r'([A-Z][a-z]+ [A-Z][a-z]+(?:,? and [A-Z][a-z]+ [A-Z][a-z]+)*)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text[:1000])  # Check first 1000 chars
            for match in matches:
                # Split by common separators
                names = re.split(r',\s*(?:and\s+)?|\s+and\s+', match)
                for name in names:
                    clean_name = re.sub(r'[^a-zA-Z\s.-]', '', name.strip())
                    if clean_name and len(clean_name.split()) >= 2:
                        authors.append(Author(name=clean_name))
                        if len(authors) >= 10:  # Limit to reasonable number
                            return authors
        
        return authors[:10]  # Return max 10 authors
    
    def _extract_abstract_from_text(self, text: str) -> str:
        """Extract abstract from text content."""
        # Common abstract patterns
        patterns = [
            r'(?:Abstract|ABSTRACT)[:\s]+(.*?)(?:\n\s*\n|Keywords?|Introduction|1\.|$)',
            r'(?:Summary|SUMMARY)[:\s]+(.*?)(?:\n\s*\n|Keywords?|Introduction|1\.|$)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text[:5000], re.DOTALL | re.IGNORECASE)
            if match:
                abstract = match.group(1).strip()
                # Clean up common formatting issues
                abstract = re.sub(r'\s+', ' ', abstract)
                abstract = abstract.replace('\n', ' ')
                if 50 <= len(abstract) <= 2000:  # Reasonable abstract length
                    return abstract
        
        return ""
    
    def _extract_keywords_from_text(self, text: str) -> List[str]:
        """Extract keywords from text content."""
        # Common keyword patterns
        patterns = [
            r'(?:Keywords?|KEY WORDS?)[:\s]+(.*?)(?:\n\s*\n|Introduction|1\.|$)',
            r'(?:Index terms?|Terms?)[:\s]+(.*?)(?:\n\s*\n|Introduction|1\.|$)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text[:3000], re.DOTALL | re.IGNORECASE)
            if match:
                keywords_text = match.group(1).strip()
                # Split by common separators
                keywords = re.split(r'[;,]|\s*—\s*|\s*-\s*', keywords_text)
                keywords = [kw.strip() for kw in keywords if kw.strip()]
                keywords = [kw for kw in keywords if 2 <= len(kw) <= 50]  # Reasonable length
                return keywords[:20]  # Limit to 20 keywords
        
        return []
    
    def _basic_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Return basic file metadata when specialized extraction fails."""
        return {
            "title": Path(file_path).stem,
            "authors": [],
            "abstract": "",
            "keywords": [],
            "total_pages": None,
            "total_words": None
        }

    def update_paper_status(
        self,
        paper_id: int,
        included_in_review: Optional[bool] = None,
        exclusion_reason: Optional[str] = None,
        quality_assessed: Optional[bool] = None,
        notes: Optional[str] = None
    ) -> ResearchPaper:
        """
        Update research paper review status for systematic literature review.

        Systematic Review Status Management:
        - Include/exclude papers based on criteria
        - Track exclusion reasons following PRISMA guidelines
        - Mark quality assessment completion
        - Maintain review audit trail

        Args:
            paper_id: Research paper to update
            included_in_review: Whether paper is included in review
            exclusion_reason: Reason for exclusion (if excluded)
            quality_assessed: Whether quality assessment is complete
            notes: Review notes and comments

        Returns:
            Updated research paper

        Raises:
            ValueError: If paper not found or status invalid
        """
        paper = self.paper_repository.get_by_id(paper_id)
        if not paper:
            raise ValueError(f"Research paper {paper_id} not found")

        # Academic Rule: Exclusion reason required if excluded
        if included_in_review is False and not exclusion_reason:
            raise ValueError("Exclusion reason required when excluding paper from review")

        # Academic Rule: Clear exclusion reason if included
        if included_in_review is True and exclusion_reason:
            exclusion_reason = None

        # Update paper status
        if included_in_review is not None:
            paper.included_in_review = included_in_review
        if exclusion_reason is not None:
            paper.exclusion_reason = exclusion_reason
        if quality_assessed is not None:
            paper.quality_assessed = quality_assessed
        if notes is not None:
            paper.notes = notes

        try:
            return self.paper_repository.update(paper)
        except Exception as e:
            raise ResearchDocumentError(f"Failed to update paper status: {str(e)}") from e

    def get_corpus_statistics(self) -> Dict[str, Any]:
        """
        Get comprehensive statistics about the research corpus.

        Academic Corpus Statistics:
        - Total papers and distribution by type
        - Publication year distribution
        - Author and journal statistics
        - Citation metrics and patterns
        - Review status distribution
        - Quality assessment progress

        Returns:
            Dictionary containing comprehensive corpus statistics
        """
        try:
            # 1. Get all papers
            all_papers = self.paper_repository.list_all()
            
            # 2. Calculate basic statistics
            stats = self._calculate_basic_corpus_statistics(all_papers)
            
            # 3. Calculate citation statistics
            stats["citation_statistics"] = self._calculate_citation_statistics(all_papers)
            
            # 4. Aggregate distributions
            distributions = self._aggregate_paper_distributions(all_papers)
            stats.update(distributions)
            
            return stats

        except Exception as e:
            raise ResearchDocumentError(f"Failed to get corpus statistics: {str(e)}") from e

    # Private helper methods for academic processing

    def _calculate_basic_corpus_statistics(
        self,
        papers: List[ResearchPaper]
    ) -> Dict[str, Any]:
        """
        Calculate basic corpus statistics.
        
        Returns:
            Dictionary with counts and review status
        """
        return {
            "total_papers": len(papers),
            "review_status": {
                "included": sum(1 for p in papers if p.included_in_review is True),
                "excluded": sum(1 for p in papers if p.included_in_review is False),
                "pending": sum(1 for p in papers if p.included_in_review is None)
            },
            "quality_assessed": sum(1 for p in papers if p.quality_assessed),
            "indexed_papers": sum(1 for p in papers if p.indexed),
            "total_size_mb": sum(p.file_size for p in papers if p.file_size) / (1024 * 1024)
        }

    def _calculate_citation_statistics(
        self,
        papers: List[ResearchPaper]
    ) -> Dict[str, Any]:
        """
        Calculate citation-related statistics.
        
        Returns:
            Dictionary with citation metrics
        """
        citation_counts = [p.citation_count for p in papers if p.citation_count]
        
        stats = {
            "total_citations": sum(p.citation_count for p in papers if p.citation_count),
            "average_citations": 0,
            "max_citations": 0,
            "papers_with_citations": 0
        }
        
        if citation_counts:
            stats["average_citations"] = int(sum(citation_counts) / len(citation_counts))
            stats["max_citations"] = max(citation_counts)
            stats["papers_with_citations"] = len(citation_counts)
        
        return stats

    def _aggregate_paper_distributions(
        self,
        papers: List[ResearchPaper]
    ) -> Dict[str, Dict[str, int]]:
        """
        Aggregate papers by various categories.
        
        Returns:
            Dictionary with methodology, study type, year, journal, file type, and author distributions
        """
        distributions: Dict[str, Dict[str, int]] = {
            "methodologies": {},
            "study_types": {},
            "publication_years": {},
            "journals": {},
            "file_types": {},
            "authors": {}
        }
        
        for paper in papers:
            # Methodology distribution
            if paper.methodology:
                distributions["methodologies"][paper.methodology] = \
                    distributions["methodologies"].get(paper.methodology, 0) + 1

            # Study type distribution
            if paper.study_type:
                distributions["study_types"][paper.study_type] = \
                    distributions["study_types"].get(paper.study_type, 0) + 1

            # Publication year distribution
            if paper.publication_year:
                year = str(paper.publication_year)
                distributions["publication_years"][year] = \
                    distributions["publication_years"].get(year, 0) + 1

            # Journal distribution
            if paper.journal and paper.journal.name:
                journal_name = paper.journal.name
                distributions["journals"][journal_name] = \
                    distributions["journals"].get(journal_name, 0) + 1

            # File type distribution
            file_type = paper.file_type or "unknown"
            distributions["file_types"][file_type] = \
                distributions["file_types"].get(file_type, 0) + 1

            # Author statistics
            if paper.authors:
                for author in paper.authors:
                    author_name = author.name
                    distributions["authors"][author_name] = \
                        distributions["authors"].get(author_name, 0) + 1
        
        # Limit to top entries for readability
        distributions["journals"] = dict(
            sorted(distributions["journals"].items(), key=lambda x: x[1], reverse=True)[:10]
        )
        distributions["authors"] = dict(
            sorted(distributions["authors"].items(), key=lambda x: x[1], reverse=True)[:20]
        )
        
        return distributions

    def _is_duplicate_doi(self, doi: str) -> bool:
        """Check if DOI already exists in corpus."""
        try:
            all_papers = self.paper_repository.list_all()
            return any(paper.doi == doi for paper in all_papers if paper.doi)
        except Exception:
            # If query fails, assume no duplicate to avoid blocking uploads
            return False

    def _group_duplicate_papers(
        self, 
        papers: List[ResearchPaper], 
        similarity_threshold: float
    ) -> List[List[ResearchPaper]]:
        """
        Group papers by similarity to identify duplicates.
        
        Uses multiple criteria:
        - Exact DOI match
        - Title similarity above threshold
        - Same title + year + first author
        
        Returns:
            List of paper groups where each group contains duplicates
        """
        groups = []
        processed_ids = set()
        
        for i, paper1 in enumerate(papers):
            if paper1.id in processed_ids:
                continue
            
            # Start new group with current paper
            current_group = [paper1]
            processed_ids.add(paper1.id)
            
            # Find similar papers
            for paper2 in papers[i+1:]:
                if paper2.id in processed_ids:
                    continue
                
                # Check duplicate criteria
                is_duplicate = False
                
                # 1. Exact DOI match
                if (paper1.doi and paper2.doi and 
                    paper1.doi.strip().lower() == paper2.doi.strip().lower()):
                    is_duplicate = True
                
                # 2. Title similarity
                elif self._calculate_title_similarity(paper1.title, paper2.title) >= similarity_threshold:
                    is_duplicate = True
                
                # 3. Same title + year + first author
                elif (paper1.title.strip().lower() == paper2.title.strip().lower() and
                      paper1.publication_year == paper2.publication_year and
                      self._same_first_author(paper1, paper2)):
                    is_duplicate = True
                
                if is_duplicate:
                    current_group.append(paper2)
                    processed_ids.add(paper2.id)
            
            groups.append(current_group)
        
        return groups

    def _build_duplicate_report(
        self,
        groups: List[List[ResearchPaper]],
        dry_run: bool
    ) -> List[Dict[str, Any]]:
        """
        Build detailed report of duplicate groups.
        
        Returns:
            List of duplicate details (limited to first 10 groups)
        """
        duplicate_details = []
        
        for group in groups:
            if len(group) <= 1:
                continue
            
            kept_paper = group[0]
            removed_papers = group[1:]
            
            detail = {
                ("would_keep_paper" if dry_run else "kept_paper"): {
                    "id": kept_paper.id,
                    "title": kept_paper.title,
                    "authors": [author.name for author in kept_paper.authors] if kept_paper.authors else []
                },
                ("would_remove_papers" if dry_run else "removed_papers"): [
                    {
                        "id": paper.id,
                        "title": paper.title,
                        "similarity": self._calculate_title_similarity(kept_paper.title, paper.title)
                    }
                    for paper in removed_papers
                ]
            }
            duplicate_details.append(detail)
        
        return duplicate_details[:10]  # Limit to first 10

    def _remove_duplicate_papers(
        self,
        groups: List[List[ResearchPaper]]
    ) -> int:
        """
        Remove duplicate papers, keeping first paper in each group.
        
        Returns:
            Number of papers successfully removed
        """
        removed_count = 0
        
        for group in groups:
            if len(group) <= 1:
                continue
            
            # Keep first paper, remove rest
            papers_to_remove = group[1:]
            
            for paper in papers_to_remove:
                try:
                    if paper.id is not None:
                        self.paper_repository.delete(paper.id)
                        removed_count += 1
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Could not remove duplicate paper {paper.id}: {str(e)}")
        
        return removed_count

    def detect_and_remove_duplicates(self, 
                                   similarity_threshold: float = 0.85,
                                   dry_run: bool = True) -> Dict[str, Any]:
        """
        Detect and optionally remove duplicate papers from the corpus.
        
        Args:
            similarity_threshold: Title similarity threshold for duplicate detection (0.0-1.0)
            dry_run: If True, only detect duplicates without removing them
            
        Returns:
            Dictionary with duplicate detection results and actions taken
        """
        try:
            # 1. Get all papers
            all_papers = self.paper_repository.list_all()
            
            if len(all_papers) < 2:
                return {
                    "success": True,
                    "message": "Not enough papers to detect duplicates",
                    "duplicates_found": 0,
                    "papers_removed": 0,
                    "total_papers": len(all_papers)
                }
            
            # 2. Group papers by duplicates
            duplicate_groups = self._group_duplicate_papers(all_papers, similarity_threshold)
            
            # 3. Count total duplicates
            total_duplicates = sum(len(group) - 1 for group in duplicate_groups if len(group) > 1)
            
            # 4. Remove duplicates if not dry run
            removed_count = 0
            if not dry_run and total_duplicates > 0:
                removed_count = self._remove_duplicate_papers(duplicate_groups)
            
            # 5. Build detailed report
            duplicate_details = self._build_duplicate_report(duplicate_groups, dry_run)
            
            # 6. Return results
            return {
                "success": True,
                "dry_run": dry_run,
                "duplicates_found": total_duplicates,
                "papers_removed": removed_count,
                "total_papers_before": len(all_papers),
                "total_papers_after": len(all_papers) - removed_count,
                "duplicate_groups": len([g for g in duplicate_groups if len(g) > 1]),
                "similarity_threshold": similarity_threshold,
                "duplicate_details": duplicate_details,
                "message": f"{'Would remove' if dry_run else 'Removed'} {total_duplicates} duplicate papers" if total_duplicates > 0 else "No duplicates found"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error during duplicate detection: {str(e)}",
                "duplicates_found": 0,
                "papers_removed": 0
            }

    def _calculate_title_similarity(self, title1: str, title2: str) -> float:
        """Calculate similarity between two paper titles."""
        if not title1 or not title2:
            return 0.0
        
        # Normalize titles
        t1 = title1.strip().lower()
        t2 = title2.strip().lower()
        
        # Exact match
        if t1 == t2:
            return 1.0
        
        # Simple word-based similarity (Jaccard similarity)
        words1 = set(t1.split())
        words2 = set(t2.split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0

    def _same_first_author(self, paper1: ResearchPaper, paper2: ResearchPaper) -> bool:
        """Check if two papers have the same first author."""
        if not paper1.authors or not paper2.authors:
            return False
        
        author1 = paper1.authors[0].name.strip().lower()
        author2 = paper2.authors[0].name.strip().lower()
        
        return author1 == author2

    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF files using academic parsing."""
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata from PDF properties
                metadata = pdf_reader.metadata if pdf_reader.metadata else {}
                
                # Extract text from first few pages for title/abstract detection
                text_content = ""
                max_pages = min(3, len(pdf_reader.pages))  # Read first 3 pages
                
                for i in range(max_pages):
                    try:
                        text_content += pdf_reader.pages[i].extract_text()
                    except Exception:
                        # Skip pages that can't be extracted
                        continue
                
                # Extract title - try PDF metadata first, then text analysis
                title = metadata.get('/Title', '').strip()
                if not title and text_content:
                    # Simple heuristic: first non-empty line is often the title
                    lines = [line.strip() for line in text_content.split('\n') if line.strip()]
                    title = lines[0] if lines else Path(file_path).stem
                
                if not title:
                    title = Path(file_path).stem
                
                return {
                    "title": title,
                    "total_pages": len(pdf_reader.pages),
                    "authors": self._extract_authors_from_text(text_content),
                    "abstract": self._extract_abstract_from_text(text_content),
                    "keywords": self._extract_keywords_from_text(text_content),
                    "total_words": len(text_content.split()) if text_content else 0
                }
                
        except ImportError:
            # Fallback if PyPDF2 not available
            return self._basic_file_metadata(file_path)
        except Exception:
            # Fallback for any PDF parsing errors
            return self._basic_file_metadata(file_path)

    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX files using academic parsing."""
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract metadata from document properties
            props = doc.core_properties
            title = props.title if props.title else ""
            
            # Extract text content from document
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # If no title from properties, extract from first paragraph
            if not title and doc.paragraphs:
                first_para = doc.paragraphs[0].text.strip()
                if first_para and len(first_para) < 200:  # Reasonable title length
                    title = first_para
            
            if not title:
                title = Path(file_path).stem
            
            return {
                "title": title,
                "authors": self._extract_authors_from_text(text_content),
                "abstract": self._extract_abstract_from_text(text_content),
                "keywords": self._extract_keywords_from_text(text_content),
                "total_words": len(text_content.split()) if text_content else 0
            }
            
        except ImportError:
            # Fallback if python-docx not available
            return self._basic_file_metadata(file_path)
        except Exception:
            # Fallback for any DOCX parsing errors
            return self._basic_file_metadata(file_path)

    def _extract_tex_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from LaTeX files using academic parsing."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract title
            title_match = re.search(r'\\title\{([^}]+)\}', content)
            title = title_match.group(1) if title_match else Path(file_path).stem
            
            # Extract authors
            author_matches = re.findall(r'\\author\{([^}]+)\}', content)
            authors = [Author(name=author.strip()) for author in author_matches]
            
            # Extract abstract
            abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', content, re.DOTALL)
            abstract = abstract_match.group(1).strip() if abstract_match else ""
            
            return {
                "title": title,
                "authors": authors,
                "abstract": abstract,
                "keywords": [],
                "total_words": len(content.split())
            }
            
        except Exception:
            return {"title": Path(file_path).stem}

    def _extract_bib_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from BibTeX files - handles multiple entries."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse all BibTeX entries
            entries = re.findall(r'@\w+\{[^,]+,\s*(.*?)\n\}', content, re.DOTALL | re.IGNORECASE)
            
            if not entries:
                return self._basic_file_metadata(file_path)
            
            # Use first entry for main metadata (for single paper upload)
            first_entry = entries[0]
            
            # Extract title from first entry
            title_match = re.search(r'title\s*=\s*[{"](.*?)[}"]', first_entry, re.IGNORECASE | re.DOTALL)
            title = title_match.group(1).strip() if title_match else Path(file_path).stem
            
            # Extract authors from first entry
            author_match = re.search(r'author\s*=\s*[{"](.*?)[}"]', first_entry, re.IGNORECASE | re.DOTALL)
            authors = []
            if author_match:
                author_text = author_match.group(1)
                author_names = [name.strip() for name in author_text.split(' and ')]
                from ..domain.models import Author
                authors = [Author(name=name) for name in author_names if name]
            
            # Extract abstract from first entry
            abstract_match = re.search(r'abstract\s*=\s*[{"](.*?)[}"]', first_entry, re.IGNORECASE | re.DOTALL)
            abstract = abstract_match.group(1).strip() if abstract_match else ""
            
            # Extract keywords from first entry
            keywords_match = re.search(r'keywords\s*=\s*[{"](.*?)[}"]', first_entry, re.IGNORECASE | re.DOTALL)
            keywords = []
            if keywords_match:
                keywords_text = keywords_match.group(1)
                keywords = [kw.strip() for kw in keywords_text.split(',') if kw.strip()]
            
            # Extract year from first entry
            year_match = re.search(r'year\s*=\s*[{"]*(\d{4})[}"]', first_entry, re.IGNORECASE)
            publication_year = int(year_match.group(1)) if year_match else None
            
            # Extract DOI from first entry
            doi_match = re.search(r'doi\s*=\s*[{"](.*?)[}"]', first_entry, re.IGNORECASE)
            doi = doi_match.group(1).strip() if doi_match else None
            
            # Extract journal from first entry
            journal_match = re.search(r'journal\s*=\s*[{"](.*?)[}"]', first_entry, re.IGNORECASE)
            journal_name = journal_match.group(1).strip() if journal_match else None
            
            metadata = {
                "title": title,
                "authors": authors,
                "abstract": abstract,
                "keywords": keywords,
                "total_words": len(content.split()),
                "publication_year": publication_year,
                "doi": doi,
                "total_entries": len(entries)  # Track how many entries in file
            }
            
            if journal_name:
                from ..domain.models import Journal
                metadata["journal"] = Journal(name=journal_name)
            
            return metadata
            
        except Exception:
            return self._basic_file_metadata(file_path)

    def _extract_ris_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from RIS (Research Information Systems) files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract title from first entry
            title_match = re.search(r'^TI\s*-\s*(.+?)$', content, re.MULTILINE)
            title = title_match.group(1).strip() if title_match else Path(file_path).stem
            
            # Extract authors from first entry
            authors = []
            author_matches = re.findall(r'^AU\s*-\s*(.+?)$', content, re.MULTILINE)
            if author_matches:
                from ..domain.models import Author
                authors = [Author(name=name.strip()) for name in author_matches if name.strip()]
            
            # Extract abstract
            abstract_match = re.search(r'^AB\s*-\s*(.+?)(?=^[A-Z]{2}\s*-|\Z)', content, 
                                     re.MULTILINE | re.DOTALL)
            abstract = abstract_match.group(1).strip() if abstract_match else ""
            
            # Extract keywords
            keywords = []
            keyword_matches = re.findall(r'^KW\s*-\s*(.+?)$', content, re.MULTILINE)
            if keyword_matches:
                for kw_line in keyword_matches:
                    keywords.extend([kw.strip() for kw in kw_line.split(';') if kw.strip()])
            
            # Extract publication year
            year_match = re.search(r'^PY\s*-\s*(\d{4})', content, re.MULTILINE)
            publication_year = int(year_match.group(1)) if year_match else None
            
            # Extract journal/venue
            journal_match = re.search(r'^JO\s*-\s*(.+?)$', content, re.MULTILINE)
            journal_name = journal_match.group(1).strip() if journal_match else None
            
            # Extract DOI
            doi_match = re.search(r'^DO\s*-\s*(.+?)$', content, re.MULTILINE)
            doi = doi_match.group(1).strip() if doi_match else None
            
            metadata = {
                "title": title,
                "authors": authors,
                "abstract": abstract,
                "keywords": keywords,
                "total_words": len(content.split()),
                "publication_year": publication_year,
                "doi": doi
            }
            
            if journal_name:
                from ..domain.models import Journal
                metadata["journal"] = Journal(name=journal_name)
            
            return metadata
            
        except Exception:
            return self._basic_file_metadata(file_path)

    def _extract_text_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic metadata from text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                "title": Path(file_path).stem,
                "total_words": len(content.split()),
                "abstract": "",
                "keywords": []
            }
        except Exception:
            return {"title": Path(file_path).stem}

    def _enhance_academic_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Enhance metadata with academic analysis."""
        enhanced = {}
        
        # Estimate sample size from abstract/title
        title = metadata.get("title", "")
        abstract = metadata.get("abstract", "")
        text = f"{title} {abstract}".lower()
        
        # Look for sample size indicators
        sample_patterns = [
            r'n\s*=\s*(\d+)',
            r'n\s*=\s*(\d+)',
            r'sample.*?(\d+)',
            r'participants.*?(\d+)',
            r'subjects.*?(\d+)'
        ]
        
        for pattern in sample_patterns:
            match = re.search(pattern, text)
            if match:
                enhanced["sample_size"] = int(match.group(1))
                break
        
        return enhanced

    def _classify_paper(self, title: str, abstract: str, keywords: List[str]) -> Dict[str, Any]:
        """Classify paper by methodology and study type using text analysis."""
        text = f"{title} {abstract} {' '.join(keywords)}".lower()
        
        classification = {}
        
        # Methodology classification
        methodology_indicators = {
            "quantitative": ["statistical", "analysis", "data", "quantitative", "numerical", "measurement"],
            "qualitative": ["qualitative", "interview", "thematic", "content analysis", "ethnographic"],
            "mixed_methods": ["mixed methods", "mixed-methods", "both qualitative and quantitative"],
            "theoretical": ["theoretical", "conceptual", "framework", "model", "theory"],
            "empirical": ["empirical", "experiment", "study", "evidence", "findings"]
        }
        
        methodology_scores = {}
        for methodology, indicators in methodology_indicators.items():
            score = sum(1 for indicator in indicators if indicator in text)
            if score > 0:
                methodology_scores[methodology] = score
        
        if methodology_scores:
            classification["methodology"] = max(methodology_scores, key=lambda x: methodology_scores.get(x, 0))
        
        # Study type classification
        study_type_indicators = {
            "experimental": ["experiment", "controlled", "randomized", "trial", "intervention"],
            "observational": ["observational", "cohort", "cross-sectional", "longitudinal"],
            "review": ["systematic review", "literature review", "meta-analysis", "review"],
            "meta-analysis": ["meta-analysis", "meta analysis", "pooled analysis"],
            "case_study": ["case study", "case series", "case report"],
            "survey": ["survey", "questionnaire", "poll", "cross-sectional survey"]
        }
        
        study_type_scores = {}
        for study_type, indicators in study_type_indicators.items():
            score = sum(1 for indicator in indicators if indicator in text)
            if score > 0:
                study_type_scores[study_type] = score
        
        if study_type_scores:
            classification["study_type"] = max(study_type_scores, key=lambda x: study_type_scores.get(x, 0))
        
        return classification

    def _analyze_paper_citations(self, paper_id: int, file_path: str) -> Dict[str, Any]:
        """Perform citation analysis for a research paper."""
        # Placeholder for comprehensive citation analysis
        # Would extract references, analyze networks, calculate metrics
        return {
            "total_references": 0,
            "citation_patterns": {},
            "reference_types": {},
            "temporal_distribution": {},
            "network_metrics": {}
        }

    def _build_academic_filters(self, filters: Dict[str, Any]) -> Dict[str, Any]:
        """Build repository filters from academic filter criteria."""
        academic_filters = {}
        
        # Map academic filters to repository filters
        field_mappings = {
            "methodology": "methodology",
            "study_type": "study_type",
            "publication_year": "publication_year",
            "included_in_review": "included_in_review",
            "quality_assessed": "quality_assessed",
            "authors": "authors",
            "tags": "tags"
        }
        
        for academic_field, repo_field in field_mappings.items():
            if academic_field in filters:
                academic_filters[repo_field] = filters[academic_field]
        
        return academic_filters

    def _sort_research_papers(
        self,
        papers: List[ResearchPaper],
        sort_by: str,
        sort_order: str
    ) -> List[ResearchPaper]:
        """Sort research papers by specified field and order."""
        reverse = sort_order.lower() == "desc"
        
        if sort_by == "publication_year":
            return sorted(papers, key=lambda p: p.publication_year or 0, reverse=reverse)
        elif sort_by == "citation_count":
            return sorted(papers, key=lambda p: p.citation_count or 0, reverse=reverse)
        elif sort_by == "title":
            return sorted(papers, key=lambda p: p.title.lower(), reverse=reverse)
        elif sort_by == "created_at":
            return sorted(papers, key=lambda p: p.created_at or datetime.min, reverse=reverse)
        elif sort_by == "upload_date":
            return sorted(papers, key=lambda p: p.upload_date or datetime.min, reverse=reverse)
        else:
            return papers

    def _targeted_academic_search(
        self,
        query: str,
        search_fields: List[str],
        filters: Optional[Dict[str, Any]],
        limit: int
    ) -> List[ResearchPaper]:
        """Perform targeted search in specific academic fields."""
        # Placeholder for field-specific academic search
        # Would implement searches in title, abstract, authors, keywords
        return self.paper_repository.search_papers(query, limit)

    def get_paper_structure(self, paper_id: int) -> Dict[str, Any]:
        """
        Extract and analyze the structure of a research paper.

        This method provides comprehensive structural analysis of academic papers
        including section detection, content organization, and academic features.

        Args:
            paper_id: ID of the paper to analyze

        Returns:
            Dictionary containing detailed structure information:
            - title, authors, basic metadata
            - sections with titles, types, word counts, page numbers
            - subsections and hierarchical structure
            - analysis metrics (citations, figures, tables)
            - content complexity and academic features

        Raises:
            ResearchDocumentError: If paper not found or structure extraction fails
        """
        try:
            # Get the paper from repository
            paper = self.paper_repository.get_by_id(paper_id)
            if not paper:
                return {"error": f"Paper with ID {paper_id} not found"}

            # Extract content from the paper file
            content = self._extract_paper_content(paper)
            if not content:
                return {"error": "Could not extract content from paper file"}

            # Build basic structure information
            structure = {
                "paper_id": paper_id,
                "title": paper.title,
                "authors": [author.name for author in paper.authors] if paper.authors else [],
                "publication_year": paper.publication_year,
                "total_pages": paper.total_pages,
                "total_words": paper.total_words or len(content.split()),
                "file_type": paper.file_type,
                "sections": [],
                "analysis": {}
            }

            # Detect and extract sections
            sections = self._extract_document_sections(content)
            structure["sections"] = sections

            # Perform content analysis
            analysis = self._analyze_document_content(content, sections)
            structure["analysis"] = analysis

            return structure

        except Exception as e:
            return {"error": f"Failed to extract paper structure: {str(e)}"}

    def _extract_paper_content(self, paper: ResearchPaper) -> str:
        """Extract content from paper file."""
        import os
        
        if not paper.file_path or not os.path.exists(paper.file_path):
            # Use available metadata if no file
            content = f"Title: {paper.title}\n\n"
            if paper.abstract:
                content += f"Abstract: {paper.abstract}\n\n"
            return content

        try:
            file_extension = Path(paper.file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(paper.file_path)
            elif file_extension in ['.txt', '.md']:
                return self._extract_from_text_file(paper.file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_word_doc(paper.file_path)
            else:
                return f"Title: {paper.title}\n\n[File type {file_extension} not supported]"
                
        except Exception as e:
            return f"Title: {paper.title}\n\n[Error extracting content: {str(e)}]"

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            # Try pymupdf first
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text.strip()
        except ImportError:
            # Fallback to PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
            except Exception as e:
                return f"[PDF extraction failed: {str(e)}]"
        except Exception as e:
            return f"[PDF extraction failed: {str(e)}]"

    def _extract_from_text_file(self, file_path: str) -> str:
        """Extract content from text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            return "[Text file encoding not supported]"
        except Exception as e:
            return f"[Text extraction failed: {str(e)}]"

    def _extract_from_word_doc(self, file_path: str) -> str:
        """Extract content from Word document."""
        try:
            import docx
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text).strip()
        except ImportError:
            return "[Word document extraction requires python-docx package]"
        except Exception as e:
            return f"[Word document extraction failed: {str(e)}]"

    def _extract_document_sections(self, content: str) -> List[Dict[str, Any]]:
        """Extract sections from document content."""
        sections = []
        
        # Academic section patterns
        section_patterns = [
            (r'^\s*(?:Abstract|ABSTRACT)\s*$', 'abstract'),
            (r'^\s*(?:I|1)\.\s*(?:Introduction|INTRODUCTION)', 'introduction'),
            (r'^\s*(?:Introduction|INTRODUCTION)\s*$', 'introduction'),
            (r'^\s*(?:Methods?|METHODS?|Methodology|METHODOLOGY)\s*$', 'methods'),
            (r'^\s*(?:II|III|2|3)\.\s*(?:Methods?|Methodology)', 'methods'),
            (r'^\s*(?:Results?|RESULTS?|Findings?|FINDINGS?)\s*$', 'results'),
            (r'^\s*(?:III|IV|V|3|4|5)\.\s*(?:Results?|Findings?)', 'results'),
            (r'^\s*(?:Discussion|DISCUSSION)\s*$', 'discussion'),
            (r'^\s*(?:IV|V|VI|4|5|6)\.\s*(?:Discussion)', 'discussion'),
            (r'^\s*(?:Conclusion|CONCLUSION|Conclusions|CONCLUSIONS)\s*$', 'conclusion'),
            (r'^\s*(?:V|VI|VII|5|6|7)\.\s*(?:Conclusion)', 'conclusion'),
            (r'^\s*(?:References?|REFERENCES?|Bibliography|BIBLIOGRAPHY)\s*$', 'references'),
        ]
        
        lines = content.split('\n')
        section_markers = []
        
        # Find section markers
        for i, line in enumerate(lines):
            for pattern, section_type in section_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    section_markers.append((i, line.strip(), section_type))
                    break
        
        # Extract sections between markers
        for j, (line_num, title, section_type) in enumerate(section_markers):
            start_line = line_num
            
            # Find end line (next section or end of document)
            if j + 1 < len(section_markers):
                end_line = section_markers[j + 1][0]
            else:
                end_line = len(lines)
            
            # Extract section content
            section_lines = lines[start_line + 1:end_line]
            section_content = '\n'.join(section_lines).strip()
            
            if section_content and len(section_content.split()) >= 5:  # Minimum content
                word_count = len(section_content.split())
                
                sections.append({
                    "title": title,
                    "type": section_type,
                    "word_count": word_count,
                    "line_start": start_line,
                    "line_end": end_line,
                    "content_preview": section_content[:200] + "..." if len(section_content) > 200 else section_content
                })
        
        return sections

    def _analyze_document_content(self, content: str, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze document content for academic features."""
        analysis: Dict[str, Any] = {}
        
        # Count citations
        citation_patterns = [
            r'\([^)]*\d{4}[^)]*\)',  # (Author, 2023)
            r'\[\d+\]',              # [1]
            r'\b[A-Z][a-z]+\s+et\s+al\.\s+\(\d{4}\)',  # Author et al. (2023)
        ]
        
        citation_count = 0
        for pattern in citation_patterns:
            matches = re.findall(pattern, content)
            citation_count += len(matches)
        
        analysis["citation_count"] = citation_count
        
        # Count figures and tables
        figure_count = len(re.findall(r'\bfigure\s+\d+\b', content.lower()))
        table_count = len(re.findall(r'\btable\s+\d+\b', content.lower()))
        
        analysis["figure_count"] = figure_count
        analysis["table_count"] = table_count
        
        # Calculate complexity score
        words = content.split()
        sentences = re.split(r'[.!?]+\s+', content)
        avg_sentence_length = len(words) / max(1, len(sentences))
        complex_words = [word for word in words if len(word) > 6]
        complexity_score = min(1.0, (avg_sentence_length / 30 + len(complex_words) / len(words)) / 2)
        
        analysis["complexity_score"] = complexity_score
        
        # Academic features
        academic_features = {
            "has_abstract": any(s["type"] == "abstract" for s in sections),
            "has_methodology": any(s["type"] == "methods" for s in sections),
            "has_results": any(s["type"] == "results" for s in sections),
            "has_discussion": any(s["type"] == "discussion" for s in sections),
            "has_conclusion": any(s["type"] == "conclusion" for s in sections),
            "has_references": any(s["type"] == "references" for s in sections),
            "sections": [s["type"] for s in sections]
        }
        
        analysis["academic_features"] = academic_features
        
        return analysis

    def upload_bibliography_batch(self, file_path: str, tags: Optional[List[str]] = None, auto_extract_metadata: bool = True) -> Dict[str, Any]:
        """
        Upload and process a bibliography file containing multiple papers.
        
        This method extracts all entries from BibTeX/RIS files and creates separate
        ResearchPaper objects for each entry, with comprehensive error tracking.
        
        Args:
            file_path: Path to bibliography file (.bib or .ris)
            tags: Optional tags to apply to all papers
            auto_extract_metadata: Whether to extract metadata automatically
            
        Returns:
            Dict with keys:
                - created_papers: List of created ResearchPaper objects
                - skipped_entries: List of dicts with {entry_num, reason, detail}
                - total_entries: Total entries found in file
                - success_count: Number of papers successfully created
                - failure_count: Number of entries that failed
                - summary: Human-readable summary
            
        Raises:
            ResearchDocumentError: If batch processing fails
        """
        try:
            file_path = str(Path(file_path).resolve())
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext not in ['.bib', '.ris']:
                raise ResearchDocumentError(f"Unsupported bibliography format: {file_ext}")
            
            if not os.path.exists(file_path):
                raise ResearchDocumentError(f"File not found: {file_path}")
            
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"📚 Starting batch upload from: {file_path}")
            
            # Extract all entries from bibliography file
            if file_ext == '.bib':
                entries_metadata = self._extract_all_bib_entries(file_path)
                logger.info(f"📄 Extracted {len(entries_metadata)} entries from BibTeX file")
                if len(entries_metadata) == 0:
                    raise ResearchDocumentError(f"EXTRACTION DEBUG: No entries found by _extract_all_bib_entries. File size: {os.path.getsize(file_path)} bytes")
            else:  # .ris
                entries_metadata = self._extract_all_ris_entries(file_path)
                logger.info(f"📄 Extracted {len(entries_metadata)} entries from RIS file")
            
            created_papers = []
            skipped_entries = []
            
            import sys
            print(f"DEBUG: Processing {len(entries_metadata)} entries", file=sys.stderr, flush=True)
            
            for i, entry_metadata in enumerate(entries_metadata):
                try:
                    # Create unique virtual file path for each entry
                    # Use UUID to make each entry's file_path unique even on re-upload
                    import uuid
                    base_name = Path(file_path).stem
                    file_dir = Path(file_path).parent
                    unique_id = str(uuid.uuid4())[:8]
                    entry_filename = f"{base_name}_entry_{i+1}_{unique_id}{file_ext}"
                    virtual_file_path = str(file_dir / entry_filename)
                    
                    # Get file stats from original file
                    file_stats = os.stat(file_path)
                    file_size = file_stats.st_size
                    
                    # Create ResearchPaper object
                    research_paper = ResearchPaper(
                        title=entry_metadata.get("title", f"Paper {i+1} from {base_name}"),
                        authors=entry_metadata.get("authors", []),
                        abstract=entry_metadata.get("abstract", ""),
                        keywords=entry_metadata.get("keywords", []),
                        journal=entry_metadata.get("journal"),
                        publication_year=entry_metadata.get("publication_year"),
                        doi=entry_metadata.get("doi"),
                        file_path=virtual_file_path,  # Unique virtual file path
                        file_type=file_ext[1:],
                        file_size=file_size,
                        total_pages=entry_metadata.get("total_pages"),
                        total_words=entry_metadata.get("total_words"),
                        tags=tags or [],
                        indexed=False,
                        quality_assessed=False,
                        included_in_review=None,
                        notes=f"Extracted from bibliography file: {Path(file_path).name}"
                    )
                    
                    # Persist research paper
                    created_paper = self.paper_repository.create(research_paper)
                    created_papers.append(created_paper)
                    print(f"DEBUG: Entry {i+1} - SUCCESS, paper created", file=sys.stderr, flush=True)
                    
                except Exception as e:
                    # Import here to avoid circular dependency
                    from ..repositories.base_repository import DuplicateEntityError
                    
                    error_type = type(e).__name__
                    error_detail = str(e)
                    print(f"DEBUG: Entry {i+1} - ERROR: {error_type}: {error_detail[:100]}", file=sys.stderr, flush=True)
                    
                    # Determine the reason for failure
                    if "already exists" in error_detail.lower() or isinstance(e, DuplicateEntityError):
                        reason = "ALREADY_EXISTS"
                    elif "parse" in error_detail.lower() or "field" in error_detail.lower():
                        reason = "PARSE_ERROR"
                    elif "validation" in error_detail.lower():
                        reason = "VALIDATION_ERROR"
                    else:
                        reason = error_type
                    
                    skipped_entries.append({
                        'entry_num': i + 1,
                        'reason': reason,
                        'error_type': error_type,
                        'detail': error_detail
                    })
                    
                    # Log to logger
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(
                        f"Entry {i+1} processing failed: {reason}",
                        extra={
                            'entry_number': i + 1,
                            'error_type': error_type,
                            'error_message': error_detail
                        }
                    )
                    continue
            
            # Build summary
            total_entries = len(entries_metadata)
            success_count = len(created_papers)
            failure_count = len(skipped_entries)
            
            summary_lines = [
                f"Processed {total_entries} entries from {Path(file_path).name}:",
                f"  ✅ Created: {success_count} papers",
                f"  ❌ Failed: {failure_count} entries"
            ]
            
            if skipped_entries:
                # Group by reason
                by_reason: Dict[str, List[int]] = {}
                for entry in skipped_entries:
                    reason: str = str(entry['reason'])  # type: ignore
                    entry_num: int = int(entry['entry_num'])  # type: ignore
                    by_reason.setdefault(reason, []).append(entry_num)
                
                summary_lines.append("  Failure breakdown:")
                for reason, entry_nums in sorted(by_reason.items()):
                    summary_lines.append(f"    - {reason}: entries {entry_nums}")
            
            if not created_papers:
                raise ResearchDocumentError("No papers could be extracted from bibliography file")
            
            return {
                'created_papers': created_papers,
                'skipped_entries': skipped_entries,
                'total_entries': total_entries,
                'success_count': success_count,
                'failure_count': failure_count,
                'summary': '\n'.join(summary_lines)
            }
            
        except Exception as e:
            raise ResearchDocumentError(f"Failed to process bibliography file: {str(e)}") from e

    def _extract_all_bib_entries(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract metadata from ALL BibTeX entries in a file."""
        import logging
        import sys
        logger = logging.getLogger(__name__)
        logger.info(f"🔍 _extract_all_bib_entries called with: {file_path}")
        
        # ALSO write to stderr for debugging
        print(f"DEBUG: _extract_all_bib_entries called with: {file_path}", file=sys.stderr, flush=True)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            logger.info(f"📖 Read {len(content)} characters from file")
            print(f"DEBUG: Read {len(content)} characters", file=sys.stderr, flush=True)
            
            # Find all complete BibTeX entries using a strict pattern first
            entry_pattern = r'(@\w+\{[^,]+,\s*.*?^\})'
            entries = re.findall(entry_pattern, content, re.DOTALL | re.MULTILINE | re.IGNORECASE)
            
            logger.info(f"🎯 Regex found {len(entries)} entries with strict pattern")
            print(f"DEBUG: Regex found {len(entries)} entries", file=sys.stderr, flush=True)

            # If strict matching fails (some exports include nested braces or unusual spacing),
            # fall back to a more tolerant split-by-entry approach.
            if not entries:
                # Split on the entry start marker while preserving the marker
                alt_entries = re.split(r'(?=@\w+\{)', content)
                # The first element may be header text before the first @ - drop it
                if alt_entries and not alt_entries[0].strip().startswith('@'):
                    alt_entries = alt_entries[1:]

                # Use non-empty chunks as entries
                entries = [e for e in alt_entries if e and e.strip()]
                # If still empty, return empty list
                if not entries:
                    return []
            
            entries_metadata = []
            
            for entry_idx, entry in enumerate(entries):
                try:
                    # Validate entry structure before parsing
                    entry_stripped = entry.strip()
                    
                    # Check 1: Must start with @
                    if not entry_stripped.startswith('@'):
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.debug(f"Entry {entry_idx + 1}: Invalid structure - does not start with @")
                        continue
                    
                    # Check 2: Must have closing brace
                    if not entry_stripped.endswith('}'):
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.debug(f"Entry {entry_idx + 1}: Malformed entry - missing closing brace")
                        continue
                    
                    # Check 3: Must have comma after entry type
                    if ',' not in entry_stripped[:100]:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.debug(f"Entry {entry_idx + 1}: Invalid structure - no comma after entry type")
                        continue
                    
                    # Extract the content between first comma and the FINAL closing brace
                    # Must handle nested braces in fields like abstract
                    # Find the opening brace after the entry type
                    open_brace_match = re.search(r'@\w+\{', entry, re.IGNORECASE)
                    if not open_brace_match:
                        continue
                    
                    # Start from after the first comma
                    comma_pos = entry.find(',', open_brace_match.end())
                    if comma_pos == -1:
                        continue
                    
                    # Find the matching closing brace by counting braces
                    brace_count = 1
                    pos = open_brace_match.end()
                    while pos < len(entry) and brace_count > 0:
                        if entry[pos] == '{':
                            brace_count += 1
                        elif entry[pos] == '}':
                            brace_count -= 1
                        pos += 1
                    
                    if brace_count != 0:
                        # Unmatched braces
                        continue
                    
                    # Extract content between comma and final closing brace
                    entry_content = entry[comma_pos + 1:pos - 1].strip()
                    
                    # Extract title - handle both {...} and "..." formats
                    title_match = re.search(r'title\s*=\s*(?:\{((?:[^{}]|(?:\{[^{}]*\}))*)\}|"([^"]*)")', entry_content, re.IGNORECASE | re.DOTALL)
                    title = (title_match.group(1) or title_match.group(2)).strip() if title_match else "Untitled"
                    
                    # Extract authors - handle both {...} and "..." formats
                    author_match = re.search(r'author\s*=\s*(?:\{((?:[^{}]|(?:\{[^{}]*\}))*)\}|"([^"]*)")', entry_content, re.IGNORECASE | re.DOTALL)
                    authors = []
                    if author_match:
                        author_text = (author_match.group(1) or author_match.group(2)) or ""
                        author_names = [name.strip() for name in author_text.split(' and ')]
                        from ..domain.models import Author
                        authors = [Author(name=name) for name in author_names if name]
                    
                    # Extract abstract - handle both {...} and "..." formats
                    abstract_match = re.search(r'abstract\s*=\s*(?:\{((?:[^{}]|(?:\{[^{}]*\}))*)\}|"([^"]*)")', entry_content, re.IGNORECASE | re.DOTALL)
                    abstract = ((abstract_match.group(1) or abstract_match.group(2)) or "").strip() if abstract_match else ""
                    
                    # Extract keywords - handle both {...} and "..." formats
                    keywords_match = re.search(r'keywords\s*=\s*(?:\{((?:[^{}]|(?:\{[^{}]*\}))*)\}|"([^"]*)")', entry_content, re.IGNORECASE | re.DOTALL)
                    keywords = []
                    if keywords_match:
                        keywords_text = (keywords_match.group(1) or keywords_match.group(2)) or ""
                        keywords = [kw.strip() for kw in keywords_text.split(',') if kw.strip()]
                    
                    # Extract year - handle both {...} and "..." formats
                    year_match = re.search(r'year\s*=\s*(?:\{(\d{4})\}|"(\d{4})")', entry_content, re.IGNORECASE)
                    publication_year = int(year_match.group(1) or year_match.group(2)) if year_match else None
                    
                    # Extract DOI - handle both {...} and "..." formats
                    doi_match = re.search(r'doi\s*=\s*(?:\{((?:[^{}]|(?:\{[^{}]*\}))*)\}|"([^"]*)")', entry_content, re.IGNORECASE)
                    doi = ((doi_match.group(1) or doi_match.group(2)) or "").strip() if doi_match else None
                    
                    # Extract journal - handle both {...} and "..." formats
                    journal_match = re.search(r'journal\s*=\s*(?:\{((?:[^{}]|(?:\{[^{}]*\}))*)\}|"([^"]*)")', entry_content, re.IGNORECASE)
                    journal_name = ((journal_match.group(1) or journal_match.group(2)) or "").strip() if journal_match else None
                    
                    metadata = {
                        "title": title,
                        "authors": authors,
                        "abstract": abstract,
                        "keywords": keywords,
                        "total_words": len(entry_content.split()),
                        "publication_year": publication_year,
                        "doi": doi
                    }
                    
                    if journal_name:
                        from ..domain.models import Journal
                        metadata["journal"] = Journal(name=journal_name)
                    
                    entries_metadata.append(metadata)
                    
                except Exception as e:
                    # Skip malformed entries but log the error
                    logger.debug(f"Skipped malformed BibTeX entry {entry_idx + 1}: {str(e)}")
                    continue
            
            logger.info(f"✅ Successfully parsed {len(entries_metadata)} entries")
            return entries_metadata
            
        except Exception as e:
            raise ResearchDocumentError(f"Failed to parse BibTeX file: {str(e)}") from e

    def _extract_all_bib_entries_pybtex_fallback(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract metadata from BibTeX file using pybtex library (fallback method).
        
        This is a more robust fallback when regex-based parsing fails.
        """
        try:
            from pybtex.database import parse_file
            from ..domain.models import Author, Journal
            import logging
            logger = logging.getLogger(__name__)
            
            entries_metadata = []
            
            try:
                # Parse file with pybtex
                bib_data = parse_file(file_path)
            except Exception as e:
                logger.warning(f"Pybtex parsing failed, falling back to regex: {str(e)}")
                return []  # Let caller fall back to regex method
            
            for entry_key, entry in bib_data.entries.items():
                try:
                    title = entry.fields.get('title', 'Untitled')
                    
                    # Parse authors
                    authors = []
                    if 'author' in entry.persons:
                        for person in entry.persons['author']:
                            author_name = str(person)
                            authors.append(Author(name=author_name))
                    
                    abstract = entry.fields.get('abstract', '')
                    keywords = []
                    if 'keywords' in entry.fields:
                        keywords = [kw.strip() for kw in entry.fields['keywords'].split(',')]
                    
                    publication_year = None
                    if 'year' in entry.fields:
                        try:
                            publication_year = int(entry.fields['year'])
                        except ValueError:
                            pass
                    
                    doi = entry.fields.get('doi')
                    journal_name = entry.fields.get('journal')
                    
                    metadata = {
                        "title": title,
                        "authors": authors,
                        "abstract": abstract,
                        "keywords": keywords,
                        "total_words": len(str(entry).split()),
                        "publication_year": publication_year,
                        "doi": doi
                    }
                    
                    if journal_name:
                        metadata["journal"] = Journal(name=journal_name)
                    
                    entries_metadata.append(metadata)
                    
                except Exception as e:
                    logger.debug(f"Failed to parse BibTeX entry {entry_key}: {str(e)}")
                    continue
            
            return entries_metadata
            
        except ImportError:
            # pybtex not installed, return empty list to let caller use regex
            import logging
            logger = logging.getLogger(__name__)
            logger.debug("pybtex not installed, skipping pybtex fallback")
            return []
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Pybtex fallback failed: {str(e)}")
            return []

    def _extract_all_ris_entries(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract metadata from ALL RIS entries in a file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split on ER (End Record) markers
            entries = re.split(r'\nER\s*-\s*\n', content)
            entries_metadata = []
            
            for entry in entries:
                if not entry.strip():
                    continue
                
                try:
                    # Extract title
                    title_match = re.search(r'^TI\s*-\s*(.+?)$', entry, re.MULTILINE)
                    title = title_match.group(1).strip() if title_match else "Untitled"
                    
                    # Extract authors
                    authors = []
                    author_matches = re.findall(r'^AU\s*-\s*(.+?)$', entry, re.MULTILINE)
                    if author_matches:
                        from ..domain.models import Author
                        authors = [Author(name=name.strip()) for name in author_matches if name.strip()]
                    
                    # Extract abstract
                    abstract_match = re.search(r'^AB\s*-\s*(.+?)(?=^[A-Z]{2}\s*-|\Z)', entry, 
                                             re.MULTILINE | re.DOTALL)
                    abstract = abstract_match.group(1).strip() if abstract_match else ""
                    
                    # Extract keywords
                    keywords = []
                    keyword_matches = re.findall(r'^KW\s*-\s*(.+?)$', entry, re.MULTILINE)
                    if keyword_matches:
                        for kw_line in keyword_matches:
                            keywords.extend([kw.strip() for kw in kw_line.split(';') if kw.strip()])
                    
                    # Extract publication year
                    year_match = re.search(r'^PY\s*-\s*(\d{4})', entry, re.MULTILINE)
                    publication_year = int(year_match.group(1)) if year_match else None
                    
                    # Extract DOI
                    doi_match = re.search(r'^DO\s*-\s*(.+?)$', entry, re.MULTILINE)
                    doi = doi_match.group(1).strip() if doi_match else None
                    
                    # Extract journal
                    journal_match = re.search(r'^JO\s*-\s*(.+?)$', entry, re.MULTILINE)
                    journal_name = journal_match.group(1).strip() if journal_match else None
                    
                    metadata = {
                        "title": title,
                        "authors": authors,
                        "abstract": abstract,
                        "keywords": keywords,
                        "total_words": len(entry.split()),
                        "publication_year": publication_year,
                        "doi": doi
                    }
                    
                    if journal_name:
                        from ..domain.models import Journal
                        metadata["journal"] = Journal(name=journal_name)
                    
                    entries_metadata.append(metadata)
                    
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.debug(f"Skipped malformed RIS entry: {str(e)}")
                    continue
            
            return entries_metadata
            
        except Exception as e:
            raise ResearchDocumentError(f"Failed to parse RIS file: {str(e)}") from e


class ResearchDocumentError(Exception):
    """
    Business logic exception for research document operations.

    Raised when academic document operations fail due to business rule violations,
    academic metadata extraction errors, or research-specific persistence failures.
    """

    def __init__(self, message: str, cause: Optional[Exception] = None):
        """
        Initialize ResearchDocumentError.

        Args:
            message: Human-readable error description
            cause: Optional underlying exception that caused this error
        """
        super().__init__(message)
        self.cause = cause